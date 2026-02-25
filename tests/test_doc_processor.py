"""utils/doc_processor.py 단위 테스트."""

import io
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt, RGBColor

from utils.doc_processor import (
    extract_doc_text,
    find_placeholders_in_doc,
    replace_placeholders_to_bytes,
)


# ───────── helpers ─────────

def _create_docx_with_text(text: str) -> Path:
    """단일 paragraph가 있는 임시 docx를 BytesIO로 만들어 경로로 저장."""
    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _docx_bytes_to_path(tmp_path: Path, content: str, filename: str = "test.docx") -> Path:
    """임시 경로에 docx 파일 생성."""
    doc = Document()
    doc.add_paragraph(content)
    out = tmp_path / filename
    doc.save(str(out))
    return out


def _docx_with_table(tmp_path: Path, cell_text: str) -> Path:
    """테이블 셀에 텍스트가 있는 docx 생성."""
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run(cell_text)
    out = tmp_path / "table.docx"
    doc.save(str(out))
    return out


# ───────── find_placeholders_in_doc ─────────

class TestFindPlaceholders:
    def test_single_placeholder(self, tmp_path):
        """단일 placeholder 탐지."""
        path = _docx_bytes_to_path(tmp_path, "내용: {{efficacy_summary}}")
        result = find_placeholders_in_doc(path)
        assert result == ["efficacy_summary"]

    def test_multiple_placeholders(self, tmp_path):
        """복수 placeholder 탐지 및 중복 제거."""
        doc = Document()
        doc.add_paragraph("{{efficacy_summary}}")
        doc.add_paragraph("{{safety_profile}}")
        doc.add_paragraph("{{efficacy_summary}}")  # 중복
        path = tmp_path / "multi.docx"
        doc.save(str(path))

        result = find_placeholders_in_doc(path)
        assert sorted(result) == ["efficacy_summary", "safety_profile"]

    def test_no_placeholder(self, tmp_path):
        """placeholder 없는 문서 → 빈 리스트."""
        path = _docx_bytes_to_path(tmp_path, "일반 텍스트입니다.")
        result = find_placeholders_in_doc(path)
        assert result == []

    def test_placeholder_in_table(self, tmp_path):
        """테이블 셀 내부 placeholder 탐지."""
        path = _docx_with_table(tmp_path, "{{dosage_administration}}")
        result = find_placeholders_in_doc(path)
        assert "dosage_administration" in result


# ───────── replace_placeholders_to_bytes ─────────

class TestReplacePlaceholders:
    def test_single_run_replacement(self, tmp_path):
        """단일 run에 있는 placeholder 치환."""
        path = _docx_bytes_to_path(tmp_path, "효능: {{efficacy_summary}}")
        replacements = {"efficacy_summary": "유효성이 뛰어남"}

        result_bytes = replace_placeholders_to_bytes(path, replacements)

        doc = Document(io.BytesIO(result_bytes))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "유효성이 뛰어남" in full_text
        assert "{{efficacy_summary}}" not in full_text

    def test_split_run_replacement(self, tmp_path):
        """run이 분리된 placeholder 치환 (핵심 엣지케이스)."""
        doc = Document()
        para = doc.add_paragraph()
        # 의도적으로 run을 쪼개서 추가
        para.add_run("{{efficacy")
        para.add_run("_sum")
        para.add_run("mary}}")
        path = tmp_path / "split_run.docx"
        doc.save(str(path))

        replacements = {"efficacy_summary": "치환된 텍스트"}
        result_bytes = replace_placeholders_to_bytes(path, replacements)

        result_doc = Document(io.BytesIO(result_bytes))
        full_text = " ".join(p.text for p in result_doc.paragraphs)
        assert "치환된 텍스트" in full_text
        assert "{{" not in full_text

    def test_unknown_placeholder_preserved(self, tmp_path):
        """replacements에 없는 placeholder는 그대로 유지."""
        path = _docx_bytes_to_path(tmp_path, "{{unknown_field}}")
        replacements = {"efficacy_summary": "내용"}

        result_bytes = replace_placeholders_to_bytes(path, replacements)

        result_doc = Document(io.BytesIO(result_bytes))
        full_text = " ".join(p.text for p in result_doc.paragraphs)
        assert "{{unknown_field}}" in full_text

    def test_table_cell_replacement(self, tmp_path):
        """테이블 셀 내부 placeholder 치환."""
        path = _docx_with_table(tmp_path, "{{safety_profile}}")
        replacements = {"safety_profile": "이상반응 정보"}

        result_bytes = replace_placeholders_to_bytes(path, replacements)

        result_doc = Document(io.BytesIO(result_bytes))
        cell_texts = []
        for table in result_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_texts.append(cell.text)

        assert any("이상반응 정보" in t for t in cell_texts)


# ───────── extract_doc_text ─────────

class TestExtractDocText:
    def test_extracts_paragraph_text(self, tmp_path):
        """본문 텍스트 추출."""
        doc = Document()
        doc.add_paragraph("첫 번째 문단")
        doc.add_paragraph("두 번째 문단")
        path = tmp_path / "extract.docx"
        doc.save(str(path))

        text = extract_doc_text(path)
        assert "첫 번째 문단" in text
        assert "두 번째 문단" in text

    def test_extracts_table_text(self, tmp_path):
        """테이블 셀 텍스트 추출."""
        path = _docx_with_table(tmp_path, "테이블 내용")
        text = extract_doc_text(path)
        assert "테이블 내용" in text

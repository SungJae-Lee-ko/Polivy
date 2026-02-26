"""DC 자료 자동화 앱 — Streamlit 메인 엔트리포인트."""

import io
import json
import logging
import os
import re
import sys
import time
from pathlib import Path

import streamlit as st
from docx import Document

# 프로젝트 루트를 sys.path에 추가
sys.path.insert(0, str(Path(__file__).parent))

from config.settings import HOSPITAL_META_PATH, PRODUCTS_JSON_PATH
from config.placeholder_queries import PLACEHOLDER_QUERIES
from utils.ai_engine import RAGEngine, CellTagMapping
from utils.doc_processor import (
    detect_taggable_cells,
    find_placeholders_in_doc,
    insert_placeholder_tags,
    replace_placeholders_to_bytes,
    TaggableCell,
    CellType,
)
from utils.pdf_loader import build_vectorstore

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ───────────────────────────── 페이지 설정 ─────────────────────────────
st.set_page_config(
    page_title="DC 자료 자동화",
    page_icon="💊",
    layout="wide",
)

st.title("💊 DC 자료 자동화")
st.caption("병원 약제위원회(DC) 상정 자료를 AI로 자동 생성합니다.")


# ───────────────────────────── 유틸 함수 ─────────────────────────────
def _load_json(path: Path) -> dict:
    """JSON 파일 로드."""
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def _save_json(path: Path, data: dict) -> None:
    """JSON 파일 저장."""
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _init_session_state() -> None:
    """Streamlit session state 초기화.

    RAG 엔진, 인덱싱 상태, 태그 에디터 상태 등을 초기화합니다.
    """
    defaults = {
        "google_api_key": os.getenv("GOOGLE_API_KEY", ""),
        "vectorstore": None,
        "rag_engine": None,
        "selected_product": None,
        "selected_hospital": None,
        "indexed_files": [],
        "indexed_chunks": 0,
        "generated_results": {},       # {질문 텍스트: 생성된 답변}
        "fillable_cells": [],          # FillableCell 목록
        "cell_fills": {},              # {(ti,ri,ci): 답변} — 최종 셀 채우기용
        # 태그 에디터 관련
        "tag_gen_cells": [],           # list[TaggableCell]
        "tag_gen_mappings": [],        # list[CellTagMapping]
        "tag_editor_active": False,    # 태그 에디터 패널 표시 여부
        "tag_editor_hospital_id": None,  # 현재 편집 중인 병원 ID
        "tag_editor_template_path": None,  # 편집 중인 템플릿 경로
        "tag_editor_is_reedit": False,  # True = 기존 태그 재편집 모드
    }
    for key, val in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = val


_init_session_state()

TEMPLATES_DIR = Path("templates")
TEMPLATES_DIR.mkdir(exist_ok=True)

# ───────────────────────────── 태그 에디터 헬퍼 함수 ─────────────────────────────


def _build_cells_from_tagged_doc(doc_path: Path) -> tuple[list[TaggableCell], dict]:
    """이미 태그된 .docx에서 {{key}} 셀 추출.

    재편집 모드에서 기존 태그 키를 selectbox 기본값으로 사용하기 위함.

    Args:
        doc_path: 태그된 .docx 파일 경로

    Returns:
        (TaggableCell 목록, {(ti, ri, ci): placeholder_key 매핑})
    """
    doc = Document(str(doc_path))
    cells = []
    key_map = {}
    placeholder_pattern = re.compile(r"\{\{(\w+)\}\}")

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            seen = set()
            for ci, cell in enumerate(row.cells):
                # 병합 셀 중복 제거
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))

                text = cell.text.strip()
                matches = placeholder_pattern.findall(text)

                if matches:
                    key = matches[0]
                    key_map[(ti, ri, ci)] = key

                    # cell_type 판별
                    if text == f"{{{{{key}}}}}":
                        cell_type = CellType.EMPTY
                        question = ""
                    else:
                        cell_type = CellType.LABEL_ONLY
                        question = text.replace(f"{{{{{key}}}}}", "").strip()

                    cells.append(TaggableCell(
                        table_index=ti,
                        row_index=ri,
                        cell_index=ci,
                        question=question or text,
                        current_text=text,
                        cell_type=cell_type,
                    ))

    return cells, key_map


def _strip_all_placeholder_tags(doc_path: Path) -> bytes:
    """재편집 저장 시 기존 {{key}} 모두 제거한 bytes 반환.

    insert_placeholder_tags() 호출 전 파일에 기록하여 이중 태그 방지.

    Args:
        doc_path: .docx 파일 경로

    Returns:
        태그 제거된 .docx 파일의 bytes
    """
    doc = Document(str(doc_path))
    placeholder_pattern = re.compile(r"\{\{\w+\}\}")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if placeholder_pattern.search(run.text):
                            run.text = placeholder_pattern.sub("", run.text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ───────────────────────────── 사이드바 ─────────────────────────────
with st.sidebar:
    st.header("⚙️ 설정")

    # API Key
    api_key_input = st.text_input(
        "Google API Key",
        value=st.session_state.google_api_key,
        type="password",
        placeholder="AIza...",
        help="Google AI Studio에서 발급한 API 키",
    )
    if api_key_input != st.session_state.google_api_key:
        st.session_state.google_api_key = api_key_input
        st.session_state.rag_engine = None

    st.divider()

    # 제품 선택
    st.subheader("1️⃣ 제품 선택")
    try:
        products_data = _load_json(PRODUCTS_JSON_PATH)
        product_list = products_data.get("products", [])
    except FileNotFoundError:
        st.error(f"products.json을 찾을 수 없습니다: {PRODUCTS_JSON_PATH}")
        product_list = []

    if product_list:
        product_names = [p["name"] for p in product_list]
        selected_product_name = st.selectbox("제품", product_names, key="product_select")
        new_product = next(p for p in product_list if p["name"] == selected_product_name)

        if new_product != st.session_state.selected_product:
            st.session_state.selected_product = new_product
            st.session_state.vectorstore = None
            st.session_state.rag_engine = None
            st.session_state.indexed_files = []
            st.session_state.indexed_chunks = 0
            st.session_state.generated_results = {}
    else:
        st.warning("등록된 제품이 없습니다.")

    st.divider()

    # 병원 선택
    st.subheader("2️⃣ 병원 선택")
    try:
        hospital_data = _load_json(HOSPITAL_META_PATH)
        hospital_list = hospital_data.get("hospitals", [])
    except FileNotFoundError:
        hospital_list = []

    real_hospitals = [h for h in hospital_list if h.get("id") != "sample_hospital"]

    if real_hospitals:
        hospital_names = [h["name"] for h in real_hospitals]
        selected_hospital_name = st.selectbox("병원", hospital_names, key="hospital_select")
        new_hospital = next(h for h in real_hospitals if h["name"] == selected_hospital_name)

        if new_hospital != st.session_state.selected_hospital:
            st.session_state.selected_hospital = new_hospital
            st.session_state.generated_results = {}
            st.session_state.auto_mapping_done = False
            st.session_state.field_mapping = []

        is_ready = new_hospital.get("mode") == "manual"
        st.caption("상태: **준비됨** ✅" if is_ready else "상태: **태그 설정 필요** ⚠️")
    else:
        st.info("등록된 병원이 없습니다.\n\n**병원 양식 관리** 탭에서 추가하세요.")
        st.session_state.selected_hospital = None


# ───────────────────────────── 메인 탭 ─────────────────────────────
tab_generate, tab_hospitals = st.tabs(["📄 문서 생성", "🏥 병원 양식 관리"])


# ══════════════════════════════════════════════════════════════════
# 탭 1: 문서 생성
# ══════════════════════════════════════════════════════════════════
with tab_generate:
    product = st.session_state.selected_product
    hospital = st.session_state.selected_hospital
    api_key = st.session_state.google_api_key

    # ── Step 1: Master Data 관리 ──
    st.header("Step 1: Master Data (기준 문서) 업로드")

    if not product:
        st.info("사이드바에서 제품을 선택하세요.")
    else:
        master_data_dir = Path(product["master_data_dir"])
        master_data_dir.mkdir(parents=True, exist_ok=True)
        existing_pdfs = sorted(master_data_dir.glob("*.pdf"))

        col1, col2 = st.columns([3, 1])
        with col1:
            if existing_pdfs:
                st.success(f"저장된 PDF: **{len(existing_pdfs)}개** — {', '.join(f.name for f in existing_pdfs)}")
            else:
                st.warning("저장된 Master Data가 없습니다. PDF를 업로드하세요.")
        with col2:
            if existing_pdfs and st.button("🗑️ 초기화", help="저장된 PDF를 모두 삭제합니다"):
                for pdf in existing_pdfs:
                    pdf.unlink()
                st.session_state.vectorstore = None
                st.session_state.rag_engine = None
                st.session_state.indexed_files = []
                st.session_state.indexed_chunks = 0
                st.rerun()

        uploaded_files = st.file_uploader(
            "PDF 추가 업로드",
            type=["pdf"],
            accept_multiple_files=True,
            key="pdf_uploader",
        )
        if uploaded_files:
            for uf in uploaded_files:
                with open(master_data_dir / uf.name, "wb") as f:
                    f.write(uf.getbuffer())
            st.success(f"{len(uploaded_files)}개 파일 저장 완료")
            st.session_state.vectorstore = None
            st.session_state.rag_engine = None

        all_pdfs = sorted(master_data_dir.glob("*.pdf"))
        if all_pdfs:
            if st.session_state.vectorstore is None:
                if not api_key:
                    st.warning("Google API 키를 입력해야 인덱싱할 수 있습니다.")
                elif st.button("🔍 인덱싱 시작", type="primary"):
                    with st.spinner(f"{len(all_pdfs)}개 PDF 인덱싱 중..."):
                        try:
                            vectorstore = build_vectorstore(
                                file_paths=[str(p) for p in all_pdfs],
                                api_key=api_key,
                            )
                            st.session_state.vectorstore = vectorstore
                            st.session_state.rag_engine = RAGEngine(vectorstore, api_key)
                            st.session_state.indexed_files = [p.name for p in all_pdfs]
                            st.session_state.indexed_chunks = vectorstore.index.ntotal
                            st.rerun()
                        except Exception as e:
                            st.error(f"인덱싱 실패: {e}")
            else:
                st.success(
                    f"✅ 인덱싱 완료 — {len(st.session_state.indexed_files)}개 문서 / "
                    f"{st.session_state.indexed_chunks}개 청크"
                )

    st.divider()

    # ── Step 2: 문서 생성 ──
    st.header("Step 2: 문서 생성")

    rag_engine: RAGEngine | None = st.session_state.rag_engine

    if not rag_engine:
        st.info("Step 1에서 Master Data를 인덱싱하세요.")
    elif not hospital:
        st.info("사이드바에서 병원을 선택하세요. 병원이 없으면 **병원 양식 관리** 탭에서 먼저 등록하세요.")
    else:
        template_path = TEMPLATES_DIR / hospital["template_file"]

        if not template_path.exists():
            st.error(f"템플릿 파일을 찾을 수 없습니다: `{template_path}`")
            st.stop()

        # 양식 분석 — {{placeholder}} 태그 탐지
        placeholders = find_placeholders_in_doc(template_path)

        if not placeholders:
            st.warning(
                "양식에서 `{{태그}}` 항목을 찾지 못했습니다. "
                "템플릿에 `{{placeholder}}` 태그가 삽입되어 있는지 확인하세요."
            )
        else:
            st.info(f"양식에서 **{len(placeholders)}개 항목** 탐지 완료")

            with st.expander("탐지된 항목 목록", expanded=False):
                for i, key in enumerate(placeholders, 1):
                    query = PLACEHOLDER_QUERIES.get(key, key)
                    st.write(f"{i}. `{{{{{key}}}}}` — {query[:60]}")

            if st.button("🤖 문서 생성", type="primary", key="gen_auto"):
                progress = st.progress(0)
                status = st.empty()
                replacements: dict[str, str] = {}

                for i, key in enumerate(placeholders):
                    query_text = PLACEHOLDER_QUERIES.get(key, key)
                    status.write(f"처리 중: **{key}** ({i+1}/{len(placeholders)})")
                    try:
                        result = rag_engine.query(
                            field_id=key,
                            custom_query=query_text,
                        )
                        replacements[key] = result.answer
                    except Exception as e:
                        replacements[key] = f"[생성 실패: {e}]"
                        logger.error("질의 실패: %s — %s", key, e)
                    progress.progress((i + 1) / len(placeholders))

                st.session_state.generated_results = replacements

                status.empty()
                progress.empty()
                st.success(f"문서 생성 완료! ({len(replacements)}개 항목)")

    st.divider()

    # ── Step 3: 결과 확인 및 다운로드 ──
    st.header("Step 3: 결과 확인 및 다운로드")

    generated: dict[str, str] = st.session_state.generated_results

    if not generated:
        st.info("Step 2에서 문서를 생성하세요.")
    else:
        st.write(f"**{len(generated)}개 항목** 생성 완료. 내용을 확인하고 필요시 수정하세요.")

        edited_results: dict[str, str] = {}
        for i, (key, answer) in enumerate(generated.items()):
            query_desc = PLACEHOLDER_QUERIES.get(key, key)[:40]
            with st.expander(f"📄 {key} — {query_desc}", expanded=False):
                edited = st.text_area(
                    "내용",
                    value=answer,
                    height=200,
                    key=f"edit_{i}",
                    label_visibility="collapsed",
                )
                edited_results[key] = edited

        st.session_state.generated_results = edited_results

        st.divider()

        if hospital:
            template_path = TEMPLATES_DIR / hospital["template_file"]
            if template_path.exists() and edited_results:
                try:
                    doc_bytes = replace_placeholders_to_bytes(
                        doc_path=template_path,
                        replacements=edited_results,
                    )
                    product_name = product["id"] if product else "unknown"
                    hospital_name = hospital["id"] if hospital else "unknown"
                    download_name = f"DC_{product_name}_{hospital_name}.docx"

                    st.download_button(
                        label="⬇️ 완성된 .docx 다운로드",
                        data=doc_bytes,
                        file_name=download_name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary",
                    )
                except Exception as e:
                    st.error(f"파일 생성 실패: {e}")


# ══════════════════════════════════════════════════════════════════
# 탭 2: 병원 양식 관리
# ══════════════════════════════════════════════════════════════════
with tab_hospitals:
    st.header("🏥 병원 양식 관리")
    st.caption("병원별 DC 신청 양식(.docx)을 등록하고 관리합니다.")

    # ── 등록된 병원 목록 ──
    try:
        h_data = _load_json(HOSPITAL_META_PATH)
        h_list = h_data.get("hospitals", [])
    except FileNotFoundError:
        h_data = {"hospitals": []}
        h_list = []

    display_list = [h for h in h_list if h.get("id") != "sample_hospital"]

    st.subheader(f"등록된 병원 ({len(display_list)}개)")

    if display_list:
        for h in display_list:
            col_name, col_status, col_file, col_edit, col_del = st.columns([3, 1.5, 2, 1.2, 1])

            col_name.write(f"**{h['name']}**")

            # 상태 배지
            is_ready = h.get("mode") == "manual"
            if is_ready:
                col_status.success("✅ 준비됨")
            else:
                col_status.warning("⚠️ 태그 필요")

            col_file.write(f"`{h['template_file']}`")

            # 재편집 버튼 (준비됨 상태일 때만)
            if is_ready:
                if col_edit.button("태그 재편집", key=f"reedit_{h['id']}"):
                    target_path = TEMPLATES_DIR / h["template_file"]
                    cells, existing_key_map = _build_cells_from_tagged_doc(target_path)
                    # synthetic mapping 생성
                    synthetic_mappings = [
                        CellTagMapping(
                            table_index=c.table_index,
                            row_index=c.row_index,
                            cell_index=c.cell_index,
                            question=c.question,
                            placeholder_key=existing_key_map.get(
                                (c.table_index, c.row_index, c.cell_index), "unknown"
                            ),
                            confidence="높음",
                        )
                        for c in cells
                    ]
                    st.session_state.tag_gen_cells = cells
                    st.session_state.tag_gen_mappings = synthetic_mappings
                    st.session_state.tag_editor_active = True
                    st.session_state.tag_editor_hospital_id = h["id"]
                    st.session_state.tag_editor_template_path = str(target_path)
                    st.session_state.tag_editor_is_reedit = True
                    st.rerun()

            # 삭제 버튼
            if col_del.button("삭제", key=f"del_{h['id']}"):
                # 병원 목록에서 제거
                h_data["hospitals"] = [x for x in h_data["hospitals"] if x["id"] != h["id"]]
                _save_json(HOSPITAL_META_PATH, h_data)
                # 템플릿 파일도 삭제
                tmpl = TEMPLATES_DIR / h["template_file"]
                if tmpl.exists():
                    tmpl.unlink()
                if st.session_state.selected_hospital and st.session_state.selected_hospital.get("id") == h["id"]:
                    st.session_state.selected_hospital = None
                st.rerun()
    else:
        st.info("등록된 병원이 없습니다. 아래에서 새 병원을 추가하세요.")

    st.divider()

    # ── 새 병원 추가 ──
    st.subheader("➕ 새 병원 추가")

    hospital_name_input = st.text_input(
        "병원 이름 *",
        placeholder="예: 서울대학교병원",
        key="new_hospital_name",
    )

    template_file_input = st.file_uploader(
        "병원 양식 파일 업로드 * (.docx)",
        type=["docx"],
        key="new_hospital_file",
        help="병원에서 요구하는 DC 신청 양식 Word 파일을 업로드하세요.",
    )

    if st.button("병원 등록", type="primary", key="register_hospital_btn"):
        if not hospital_name_input:
            st.error("병원 이름을 입력하세요.")
        elif not template_file_input:
            st.error("양식 파일을 업로드하세요.")
        else:
            # 병원 ID 생성 (이름 → 영문 소문자 + 타임스탬프)
            hospital_id = re.sub(r"[^a-zA-Z0-9가-힣]", "_", hospital_name_input).lower()
            hospital_id = f"{hospital_id}_{int(time.time())}"

            # 파일명: 병원ID.docx
            template_filename = f"{hospital_id}.docx"
            save_path = TEMPLATES_DIR / template_filename

            # 파일 저장
            with open(save_path, "wb") as f:
                f.write(template_file_input.getbuffer())

            # 자동 감지: 기존 {{placeholder}} 태그 확인
            detected_placeholders = find_placeholders_in_doc(save_path)
            detected_mode = "manual" if detected_placeholders else "needs_tagging"

            # hospital_meta.json 업데이트
            new_entry = {
                "id": hospital_id,
                "name": hospital_name_input,
                "template_file": template_filename,
                "format": "docx",
                "mode": detected_mode,
                "field_mapping": None,
            }
            h_data["hospitals"].append(new_entry)
            _save_json(HOSPITAL_META_PATH, h_data)

            if detected_mode == "manual":
                # 태그가 이미 있음 — 완료
                st.success(
                    f"✅ **{hospital_name_input}** 등록 완료! "
                    f"{{{{placeholder}}}} 태그 {len(detected_placeholders)}개 발견. "
                    f"사이드바에서 선택하여 사용하세요."
                )
                # 폼 리셋
                st.session_state.new_hospital_name = ""
                st.session_state.new_hospital_file = None
                st.rerun()
            else:
                # 태그 없음 — 태그 에디터 자동 활성화
                with st.spinner("양식 분석 중..."):
                    cells = detect_taggable_cells(save_path)
                st.session_state.tag_gen_cells = cells
                st.session_state.tag_gen_mappings = []
                st.session_state.tag_editor_active = True
                st.session_state.tag_editor_hospital_id = hospital_id
                st.session_state.tag_editor_template_path = str(save_path)
                st.session_state.tag_editor_is_reedit = False
                st.rerun()

    st.divider()

    # ── 인라인 태그 에디터 패널 ──
    if st.session_state.tag_editor_active:
        st.divider()

        # 현재 편집 중인 병원 찾기
        editor_hospital = next(
            (h for h in display_list if h["id"] == st.session_state.tag_editor_hospital_id),
            None,
        )

        if editor_hospital is None:
            st.error("병원을 찾을 수 없습니다.")
        else:
            # 헤더
            if st.session_state.tag_editor_is_reedit:
                st.subheader(f"🏷️ 태그 재편집: {editor_hospital['name']}")
            else:
                st.subheader(f"🏷️ 태그 설정: {editor_hospital['name']}")
                st.info("업로드된 양식에서 태그를 찾지 못했습니다. 아래에서 각 셀에 적합한 항목을 지정하세요.")

            tag_cells: list[TaggableCell] = st.session_state.tag_gen_cells

            if not tag_cells:
                st.warning("태그 가능한 셀을 찾지 못했습니다. 양식에 테이블이 없거나 이미 모두 채워져 있을 수 있습니다.")
            else:
                # AI 자동 채우기 버튼 (선택적)
                tag_api_key = st.session_state.google_api_key
                col_ai, col_spacer = st.columns([2, 5])
                with col_ai:
                    if not tag_api_key:
                        st.caption("AI 자동 채우기를 사용하려면 사이드바에서 API 키를 입력하세요.")
                    else:
                        if st.button(
                            "🤖 AI 자동 채우기",
                            key="ai_fill_tags",
                            help="AI가 각 셀에 적합한 태그를 분석하여 제안합니다.",
                        ):
                            with st.spinner("AI가 태그를 분석 중..."):
                                tag_engine = RAGEngine(vectorstore=None, api_key=tag_api_key)
                                mappings = tag_engine.generate_cell_tags(
                                    cells=tag_cells,
                                    placeholder_queries=PLACEHOLDER_QUERIES,
                                )
                                st.session_state.tag_gen_mappings = mappings
                            st.rerun()

                # selectbox 옵션 및 AI 추천 매핑
                all_keys = ["(미지정)"] + sorted(PLACEHOLDER_QUERIES.keys())
                ai_lookup = {}
                for m in st.session_state.tag_gen_mappings:
                    if m.placeholder_key not in ("unknown", ""):
                        ai_lookup[(m.table_index, m.row_index, m.cell_index)] = m.placeholder_key

                # 각 셀 편집
                edited_assignments: list[tuple[TaggableCell, str]] = []
                st.write("**셀별 태그 지정**")

                for i, cell in enumerate(tag_cells):
                    coord = (cell.table_index, cell.row_index, cell.cell_index)
                    col_loc, col_label, col_select, col_preview = st.columns([1.5, 3, 2.5, 3])

                    with col_loc:
                        st.caption(f"T{cell.table_index}R{cell.row_index}C{cell.cell_index}")
                    with col_label:
                        st.write(cell.question[:50] if cell.question else "(라벨 없음)")
                    with col_select:
                        ai_suggestion = ai_lookup.get(coord, "")
                        default_idx = all_keys.index(ai_suggestion) if ai_suggestion in all_keys else 0

                        selected = st.selectbox(
                            "태그",
                            options=all_keys,
                            index=default_idx,
                            key=f"tag_sel_{i}",
                            label_visibility="collapsed",
                        )
                    with col_preview:
                        if selected != "(미지정)":
                            if cell.cell_type == CellType.LABEL_ONLY:
                                preview = f"`{cell.current_text.rstrip()} {{{{{selected}}}}}`"
                            else:
                                preview = f"`{{{{{selected}}}}}`"
                            st.caption(preview)
                        else:
                            st.caption("—")

                    if selected != "(미지정)":
                        edited_assignments.append((cell, selected))

                st.divider()

                # 하단 버튼
                col_save, col_skip, col_preview_dl, col_cancel = st.columns([1.5, 1.5, 1.5, 1])

                with col_preview_dl:
                    if edited_assignments:
                        preview_bytes = insert_placeholder_tags(
                            st.session_state.tag_editor_template_path,
                            edited_assignments,
                        )
                        st.download_button(
                            label="미리보기 다운로드",
                            data=preview_bytes,
                            file_name=f"tagged_preview_{editor_hospital['template_file']}",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="tag_preview_dl",
                        )

                with col_save:
                    if st.button(
                        "✅ 저장",
                        type="primary",
                        key="confirm_tags",
                        disabled=len(edited_assignments) == 0,
                    ):
                        with st.spinner("태그 삽입 중..."):
                            if st.session_state.tag_editor_is_reedit:
                                # 재편집: 기존 태그 제거 후 삽입
                                cleaned = _strip_all_placeholder_tags(
                                    Path(st.session_state.tag_editor_template_path)
                                )
                                with open(st.session_state.tag_editor_template_path, "wb") as f:
                                    f.write(cleaned)

                            tagged_bytes = insert_placeholder_tags(
                                st.session_state.tag_editor_template_path,
                                edited_assignments,
                            )

                        # 파일 저장
                        with open(st.session_state.tag_editor_template_path, "wb") as f:
                            f.write(tagged_bytes)

                        # hospital_meta.json 업데이트: mode → "manual"
                        target_id = st.session_state.tag_editor_hospital_id
                        for h in h_data["hospitals"]:
                            if h["id"] == target_id:
                                h["mode"] = "manual"
                                break
                        _save_json(HOSPITAL_META_PATH, h_data)

                        # session state 정리
                        st.session_state.tag_editor_active = False
                        st.session_state.tag_editor_hospital_id = None
                        st.session_state.tag_editor_template_path = None
                        st.session_state.tag_editor_is_reedit = False
                        st.session_state.tag_gen_cells = []
                        st.session_state.tag_gen_mappings = []

                        # 선택된 병원의 mode도 업데이트
                        if (st.session_state.selected_hospital
                                and st.session_state.selected_hospital.get("id") == target_id):
                            st.session_state.selected_hospital["mode"] = "manual"

                        st.success("✅ 태그 저장 완료! '문서 생성' 탭에서 사용 가능합니다.")
                        st.rerun()

                with col_skip:
                    if st.button(
                        "⏭️ AI 자동 분석 후 저장",
                        key="skip_tags",
                        help="AI가 자동으로 태그를 분석하여 저장합니다.",
                    ):
                        tag_api_key = st.session_state.google_api_key
                        if not tag_api_key:
                            st.error("API 키를 입력해야 AI 자동 분석을 진행할 수 있습니다. 사이드바에서 입력하세요.")
                        else:
                            with st.spinner("🤖 AI가 태그를 자동으로 분석 중... (이 과정은 몇 초 걸릴 수 있습니다)"):
                                # AI 태그 분석
                                tag_engine = RAGEngine(vectorstore=None, api_key=tag_api_key)
                                auto_mappings = tag_engine.generate_cell_tags(
                                    cells=tag_cells,
                                    placeholder_queries=PLACEHOLDER_QUERIES,
                                )

                                # AI 분석 결과로 자동 태그 삽입
                                auto_assignments = [
                                    (c, m.placeholder_key)
                                    for c, m in zip(tag_cells, auto_mappings)
                                    if m.placeholder_key not in ("unknown", "")
                                ]

                            if auto_assignments:
                                with st.spinner("태그를 삽입 중..."):
                                    if st.session_state.tag_editor_is_reedit:
                                        # 재편집: 기존 태그 제거 후 삽입
                                        cleaned = _strip_all_placeholder_tags(
                                            Path(st.session_state.tag_editor_template_path)
                                        )
                                        with open(st.session_state.tag_editor_template_path, "wb") as f:
                                            f.write(cleaned)

                                    tagged_bytes = insert_placeholder_tags(
                                        st.session_state.tag_editor_template_path,
                                        auto_assignments,
                                    )

                                # 파일 저장
                                with open(st.session_state.tag_editor_template_path, "wb") as f:
                                    f.write(tagged_bytes)

                                # hospital_meta.json 업데이트: mode → "manual"
                                target_id = st.session_state.tag_editor_hospital_id
                                for h in h_data["hospitals"]:
                                    if h["id"] == target_id:
                                        h["mode"] = "manual"
                                        break
                                _save_json(HOSPITAL_META_PATH, h_data)

                                # session state 정리
                                st.session_state.tag_editor_active = False
                                st.session_state.tag_editor_hospital_id = None
                                st.session_state.tag_editor_template_path = None
                                st.session_state.tag_editor_is_reedit = False
                                st.session_state.tag_gen_cells = []
                                st.session_state.tag_gen_mappings = []

                                # 선택된 병원의 mode도 업데이트
                                if (st.session_state.selected_hospital
                                        and st.session_state.selected_hospital.get("id") == target_id):
                                    st.session_state.selected_hospital["mode"] = "manual"

                                st.success(f"✅ AI가 {len(auto_assignments)}개 셀에 태그를 자동 삽입했습니다! 이제 '문서 생성' 탭에서 사용 가능합니다.")
                                st.rerun()
                            else:
                                st.warning("AI가 분석 가능한 태그를 찾지 못했습니다. 수동으로 태그를 설정해주세요.")

                with col_cancel:
                    if st.button("✖️ 취소", key="cancel_tag_editor"):
                        # 새 등록인 경우: 병원 + 파일 삭제
                        if not st.session_state.tag_editor_is_reedit:
                            target_id = st.session_state.tag_editor_hospital_id
                            h_data["hospitals"] = [
                                x for x in h_data["hospitals"] if x["id"] != target_id
                            ]
                            _save_json(HOSPITAL_META_PATH, h_data)
                            tmpl = TEMPLATES_DIR / editor_hospital["template_file"]
                            if tmpl.exists():
                                tmpl.unlink()

                        # 태그 에디터 초기화
                        st.session_state.tag_editor_active = False
                        st.session_state.tag_editor_hospital_id = None
                        st.session_state.tag_editor_template_path = None
                        st.session_state.tag_editor_is_reedit = False
                        st.session_state.tag_gen_cells = []
                        st.session_state.tag_gen_mappings = []
                        st.rerun()

"""Word(.docx) 템플릿 처리.

세 가지 모드:
1. placeholder 모드: {{key}} 패턴을 찾아 치환 (manual 모드)
2. 테이블 셀 모드: 빈 셀을 자동 탐지하여 채우기 (auto 모드)
3. 자동 태그 생성: 빈 셀/라벨 셀을 탐지하여 {{placeholder}} 태그를 자동 삽입
"""

import io
import logging
import re
from dataclasses import dataclass
from enum import Enum
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

from config.settings import PLACEHOLDER_PATTERN

logger = logging.getLogger(__name__)


@dataclass
class RunFormat:
    """Run의 서식 정보를 저장하는 데이터 클래스."""

    bold: bool | None = None
    italic: bool | None = None
    underline: bool | None = None
    font_name: str | None = None
    font_size: Pt | None = None
    font_color: RGBColor | None = None


def _capture_run_format(run) -> RunFormat:
    """Run의 현재 서식을 캡처."""
    color = None
    try:
        if run.font.color and run.font.color.type is not None:
            color = run.font.color.rgb
    except Exception:
        pass

    return RunFormat(
        bold=run.bold,
        italic=run.italic,
        underline=run.underline,
        font_name=run.font.name,
        font_size=run.font.size,
        font_color=color,
    )


def _apply_run_format(run, fmt: RunFormat) -> None:
    """Run에 저장된 서식을 적용."""
    if fmt.bold is not None:
        run.bold = fmt.bold
    if fmt.italic is not None:
        run.italic = fmt.italic
    if fmt.underline is not None:
        run.underline = fmt.underline
    if fmt.font_name:
        run.font.name = fmt.font_name
    if fmt.font_size:
        run.font.size = fmt.font_size
    if fmt.font_color:
        run.font.color.rgb = fmt.font_color


def _replace_in_paragraph(paragraph: Paragraph, replacements: dict[str, str]) -> bool:
    """단일 paragraph 내 모든 placeholder를 치환.

    Run이 분리된 경우도 처리: paragraph의 모든 run을 합산한 후
    치환된 텍스트를 단일 run으로 재구성.

    Args:
        paragraph: python-docx Paragraph 객체.
        replacements: {placeholder_key: 치환할 텍스트} 딕셔너리.

    Returns:
        치환이 발생했으면 True.
    """
    if not paragraph.runs:
        return False

    full_text = "".join(run.text for run in paragraph.runs)

    # 이 paragraph에 치환할 placeholder가 있는지 확인
    found_keys = PLACEHOLDER_PATTERN.findall(full_text)
    if not found_keys:
        return False

    # 치환 적용
    new_text = full_text
    changed = False
    for key in found_keys:
        if key in replacements:
            new_text = new_text.replace(f"{{{{{key}}}}}", replacements[key])
            changed = True
        # replacements에 없는 key는 그대로 둠

    if not changed:
        return False

    # 첫 번째 run의 서식 저장
    first_run_fmt = _capture_run_format(paragraph.runs[0])

    # 모든 run의 XML 요소 삭제
    p_elem = paragraph._p
    for run in paragraph.runs:
        p_elem.remove(run._r)

    # 새 run 추가 후 서식 복원
    new_run = paragraph.add_run(new_text)
    _apply_run_format(new_run, first_run_fmt)

    return True


def find_placeholders_in_doc(doc_path: str | Path) -> list[str]:
    """Word 문서 전체에서 {{placeholder}} 패턴을 탐색.

    본문 paragraph와 테이블 셀 내부를 모두 탐색.

    Args:
        doc_path: .docx 파일 경로.

    Returns:
        발견된 placeholder 키 목록 (중복 제거, 정렬).
    """
    doc = Document(str(doc_path))
    keys: set[str] = set()

    def _scan_paragraph(para: Paragraph) -> None:
        """Paragraph에서 placeholder를 찾아 keys 집합에 추가."""
        full_text = "".join(run.text for run in para.runs)
        found = PLACEHOLDER_PATTERN.findall(full_text)
        keys.update(found)

    # 본문 paragraph
    for para in doc.paragraphs:
        _scan_paragraph(para)

    # 테이블 셀
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _scan_paragraph(para)

    return sorted(keys)


def extract_doc_text(doc_path: str | Path) -> str:
    """Word 문서 전체 텍스트를 하나의 문자열로 추출 (auto 모드용).

    Args:
        doc_path: .docx 파일 경로.

    Returns:
        문서 전체 텍스트. 본문 + 테이블 내용 포함.
    """
    doc = Document(str(doc_path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    parts.append(cell_text)

    return "\n".join(parts)


def replace_placeholders_in_doc(
    doc_path: str | Path,
    replacements: dict[str, str],
    output_path: str | Path,
) -> int:
    """Word 문서의 placeholder를 일괄 치환하여 새 파일로 저장.

    Args:
        doc_path: 원본 .docx 템플릿 파일 경로.
        replacements: {placeholder_key: 치환할 텍스트} 딕셔너리.
        output_path: 결과 파일을 저장할 경로.

    Returns:
        실제로 치환된 placeholder 수.
    """
    doc = Document(str(doc_path))
    replaced_count = 0

    # 본문 paragraph 처리
    for para in doc.paragraphs:
        if _replace_in_paragraph(para, replacements):
            replaced_count += 1

    # 테이블 셀 처리
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if _replace_in_paragraph(para, replacements):
                        replaced_count += 1

    doc.save(str(output_path))
    logger.info(
        "%d개 항목 치환 완료: %s → %s",
        replaced_count,
        Path(doc_path).name,
        Path(output_path).name,
    )
    return replaced_count


def replace_placeholders_to_bytes(
    doc_path: str | Path,
    replacements: dict[str, str],
) -> bytes:
    """Word 문서의 placeholder를 치환하고 바이트로 반환 (디스크 저장 없이 다운로드용).

    Args:
        doc_path: 원본 .docx 템플릿 파일 경로.
        replacements: {placeholder_key: 치환할 텍스트} 딕셔너리.

    Returns:
        완성된 .docx 파일의 바이트 데이터.
    """
    import io

    doc = Document(str(doc_path))

    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ═══════════════════════════════════════════════════════════════
# 테이블 셀 기반 채우기 (auto 모드)
# ═══════════════════════════════════════════════════════════════

@dataclass
class FillableCell:
    """채워야 할 빈 셀 정보."""

    table_index: int
    row_index: int
    cell_index: int
    question: str       # 인접 라벨 셀의 텍스트 (RAG 쿼리로 사용)
    current_text: str   # 현재 셀 내용 (보통 빈 문자열)


def analyze_template_tables(doc_path: str | Path) -> list[FillableCell]:
    """템플릿 .docx의 테이블을 분석하여 채워야 할 빈 셀 목록을 반환.

    로직:
    - 각 테이블의 각 행을 순회
    - 빈 셀과 비어있지 않은 셀(라벨)을 구분
    - 라벨 셀의 텍스트를 해당 빈 셀의 '질문'으로 매핑

    Args:
        doc_path: .docx 파일 경로.

    Returns:
        FillableCell 목록 (채울 항목들).
    """
    doc = Document(str(doc_path))
    fillable: list[FillableCell] = []

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            cells_text = [cell.text.strip() for cell in row.cells]

            # 병합된 셀 중복 제거 (python-docx에서 병합 셀은 동일 객체를 반복 반환)
            seen_cells: set[int] = set()
            unique_indices: list[int] = []
            for ci, cell in enumerate(row.cells):
                cell_id = id(cell._tc)
                if cell_id not in seen_cells:
                    seen_cells.add(cell_id)
                    unique_indices.append(ci)

            # 라벨(비어있지 않은 셀)과 빈 셀 분리
            labels: list[str] = []
            empty_indices: list[int] = []

            for ci in unique_indices:
                text = cells_text[ci]
                if text:
                    labels.append(text)
                else:
                    empty_indices.append(ci)

            # 빈 셀이 있고 라벨도 있을 때 → 채울 대상
            if empty_indices and labels:
                question = " / ".join(labels)
                for ci in empty_indices:
                    fillable.append(
                        FillableCell(
                            table_index=ti,
                            row_index=ri,
                            cell_index=ci,
                            question=question,
                            current_text=cells_text[ci],
                        )
                    )

    logger.info(
        "%s: %d개 빈 셀 탐지 (테이블 %d개)",
        Path(doc_path).name,
        len(fillable),
        len(doc.tables),
    )
    return fillable


def fill_cells_to_bytes(
    doc_path: str | Path,
    fills: dict[tuple[int, int, int], str],
) -> bytes:
    """특정 테이블 셀 좌표에 텍스트를 삽입하고 바이트로 반환.

    Args:
        doc_path: 원본 .docx 템플릿 파일 경로.
        fills: {(table_idx, row_idx, cell_idx): "삽입할 텍스트"} 딕셔너리.

    Returns:
        완성된 .docx 파일의 바이트 데이터.
    """
    doc = Document(str(doc_path))

    for (ti, ri, ci), text in fills.items():
        try:
            cell = doc.tables[ti].rows[ri].cells[ci]
            # 기존 paragraph가 있으면 첫 번째에 텍스트 설정
            if cell.paragraphs:
                para = cell.paragraphs[0]
                # 기존 run이 있으면 서식 보존
                if para.runs:
                    fmt = _capture_run_format(para.runs[0])
                    # 기존 run 제거
                    for run in para.runs:
                        para._p.remove(run._r)
                    new_run = para.add_run(text)
                    _apply_run_format(new_run, fmt)
                else:
                    para.add_run(text)
            else:
                cell.text = text
        except (IndexError, AttributeError) as e:
            logger.warning("셀 채우기 실패 T%dR%dC%d: %s", ti, ri, ci, e)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ═══════════════════════════════════════════════════════════════
# 자동 태그 생성 (빈 셀/라벨 셀 탐지 → {{placeholder}} 삽입)
# ═══════════════════════════════════════════════════════════════

class CellType(str, Enum):
    """태그 삽입 대상 셀의 유형."""

    EMPTY = "empty"            # 빈 셀 → {{key}} 직접 삽입
    LABEL_ONLY = "label_only"  # "한글:" 같은 라벨만 있는 셀 → "한글: {{key}}" 추가


@dataclass
class TaggableCell:
    """자동 태그 삽입 후보 셀."""

    table_index: int
    row_index: int
    cell_index: int
    question: str       # 같은 행의 라벨 텍스트 (LLM 매핑에 사용)
    current_text: str   # 현재 셀 내용 ("" 또는 "한글:" 같은 라벨)
    cell_type: CellType


_LABEL_SUFFIX_RE = re.compile(r"^.+[：:)\)\]]\s*$")


def _is_label_only_cell(text: str) -> bool:
    """셀 텍스트가 라벨만 있는 셀인지 판별 (예: '한글:', '판매회사：')."""
    stripped = text.strip()
    if not stripped:
        return False
    return bool(_LABEL_SUFFIX_RE.match(stripped))


def detect_taggable_cells(doc_path: str | Path) -> list[TaggableCell]:
    """템플릿 .docx를 분석하여 자동 태그 삽입 후보 셀 목록을 반환.

    analyze_template_tables()의 확장판.
    빈 셀(EMPTY)뿐만 아니라 라벨만 있는 셀(LABEL_ONLY)도 탐지.

    Args:
        doc_path: .docx 파일 경로.

    Returns:
        TaggableCell 목록.
    """
    doc = Document(str(doc_path))
    taggable: list[TaggableCell] = []

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            cells_text = [cell.text.strip() for cell in row.cells]

            # 병합된 셀 중복 제거
            seen_cells: set[int] = set()
            unique_indices: list[int] = []
            for ci, cell in enumerate(row.cells):
                cell_id = id(cell._tc)
                if cell_id not in seen_cells:
                    seen_cells.add(cell_id)
                    unique_indices.append(ci)

            # 빈 셀과 비어있지 않은 셀 분류
            empty_indices: list[int] = []
            nonempty_indices: list[int] = []

            for ci in unique_indices:
                if cells_text[ci]:
                    nonempty_indices.append(ci)
                else:
                    empty_indices.append(ci)

            # 패턴 1: 빈 셀이 있는 경우 (라벨과 무관하게 처리)
            if empty_indices:
                question = " / ".join(cells_text[ci] for ci in nonempty_indices) if nonempty_indices else ""
                for ci in empty_indices:
                    taggable.append(
                        TaggableCell(
                            table_index=ti,
                            row_index=ri,
                            cell_index=ci,
                            question=question,
                            current_text="",
                            cell_type=CellType.EMPTY,
                        )
                    )

            # 패턴 2: 라벨만 있는 셀 (독립적으로 실행 — 빈 셀 여부와 무관)
            label_only_indices = [
                ci for ci in nonempty_indices
                if _is_label_only_cell(cells_text[ci])
            ]
            for ci in label_only_indices:
                taggable.append(
                    TaggableCell(
                        table_index=ti,
                        row_index=ri,
                        cell_index=ci,
                        question=cells_text[ci],
                        current_text=cells_text[ci],
                        cell_type=CellType.LABEL_ONLY,
                    )
                )

    logger.info(
        "%s: %d개 태그 후보 셀 탐지 (테이블 %d개)",
        Path(doc_path).name,
        len(taggable),
        len(doc.tables),
    )
    return taggable


def insert_placeholder_tags(
    doc_path: str | Path,
    tag_assignments: list[tuple[TaggableCell, str]],
) -> bytes:
    """TaggableCell 목록에 {{placeholder}} 태그를 삽입하여 바이트로 반환.

    EMPTY 셀: "{{key}}" 직접 삽입.
    LABEL_ONLY 셀: "기존텍스트 {{key}}" 형태로 텍스트 뒤에 추가.

    Args:
        doc_path: 원본 .docx 파일 경로.
        tag_assignments: [(TaggableCell, placeholder_key), ...] 리스트.
                         placeholder_key가 빈 문자열이면 해당 셀 건너뜀.

    Returns:
        태그가 삽입된 .docx 파일의 바이트 데이터.
    """
    fills: dict[tuple[int, int, int], str] = {}

    for cell, key in tag_assignments:
        if not key:
            continue
        tag = f"{{{{{key}}}}}"  # → {{key}}
        coord = (cell.table_index, cell.row_index, cell.cell_index)

        if cell.cell_type == CellType.EMPTY:
            fills[coord] = tag
        elif cell.cell_type == CellType.LABEL_ONLY:
            # 원본 텍스트에서 공백을 보존하고 태그 추가
            original = cell.current_text
            # 세미콜론으로 끝나면 공백 추가 후 태그 삽입
            if original.endswith(":") or original.endswith("："):
                fills[coord] = f"{original} {tag}"
            else:
                # 다른 경우도 공백 추가 후 태그 삽입
                fills[coord] = f"{original} {tag}"

    return fill_cells_to_bytes(doc_path, fills)
